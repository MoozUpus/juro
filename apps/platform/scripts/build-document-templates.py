#!/usr/bin/env python3
"""Create JURO runtime OOXML bases from the approved receipt template."""

from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image


def add_field(paragraph, instruction: str, display: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_footer(doc: Document, logo_path: Path) -> None:
    for section in doc.sections:
        section.footer_distance = Inches(0.35)
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        picture = paragraph.add_run().add_picture(str(logo_path), width=Inches(0.15))
        picture._inline.docPr.set("descr", "JURO")
        run = paragraph.add_run("  Создано в JURO  ·  Страница ")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = None
        add_field(paragraph, " PAGE ", "1")
        run = paragraph.add_run(" из ")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        add_field(paragraph, " NUMPAGES ", "1")
        for footer_run in paragraph.runs:
            footer_run.font.name = "Arial"
            footer_run.font.size = Pt(8)


def main() -> None:
    if len(sys.argv) != 4:
        raise SystemExit("usage: build-document-templates.py SOURCE_DOCX JURO_MARK OUTPUT_DIR")
    source, mark, output_dir = map(Path, sys.argv[1:])
    output_dir.mkdir(parents=True, exist_ok=True)
    small_mark = output_dir / "juro-mark-footer.png"
    with Image.open(mark) as image:
        image.thumbnail((96, 96), Image.Resampling.LANCZOS)
        image.save(small_mark, optimize=True)

    for name in ("receipt-ru.docx", "receipt-uz-cyrl.docx"):
        document = Document(source)
        add_footer(document, small_mark)
        properties = document.core_properties
        properties.title = "JURO receipt runtime template"
        properties.subject = "Approved JURO receipt OOXML base"
        properties.author = "JURO"
        properties.last_modified_by = "JURO"
        properties.keywords = "JURO, receipt, runtime template, Russian, Uzbek Cyrillic"
        document.save(output_dir / name)


if __name__ == "__main__":
    main()
